import re

from io import BytesIO
from pypdf import PdfReader
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
# from docx2pdf import convert

from scopetoproposal.models import Image



permit = '\nA permit is requred for projects of value in excess of R60,000,000 (Sixty Million Rand).\n'

inspection_day = ''

extra_directors = ''

# formatting = (run.bold, run.italic, run.font.color.rgb, run.font.name, run.font.size)
FOOTER_LEGAL = [
    {'text': f'Directors: L.P. Dicks, A. Spanos {extra_directors}  Co Reg: 2007/007252/07  Tel: ', 'formatting': 'normal'},
    {'text': '{COMPANY_TEL_NUMBER}', 'formatting' : 'red'},
    {'text': '\nSubsidiary of ComPrac Holdings (Pty) Ltd', 'formatting': 'normal'},
    {'text': '\n{COMPANY_NAME}', 'formatting':  'red'},
    {'text': ' is compliant with ISO 9001 and BS OHSAS 18001\nComPrac Holdings (Pty) Ltd is an Approved Inspection Authority and registered with HWSETA\nCorporate Membership:   SAIOSH & SAFCEC', 'formatting': 'normal'},
    {'text': '\nIMT-002', 'formatting': 'small'},
]




# TEMP:
# placeholders = {
#     '{CLIENT_NAME}': 'Harbour Terrace Body Corporate',
#     '{CLIENT_ADDRESS}': '9 Boundary Road, Green Point, Cape Town'.replace(', ', '\n'),
#     '{DATE}': '',
#     '{EMAIL}': 'foxmoss@isoft.co.za',
#     '{AGENT_NAME}': 'Maon Moss',
#     '{SITE_NAME}': 'Harbour Terrace',
#     '{PROJECT_NAME}' : 'Harbour Terrace',
#     '{QP_ADDRESSEE}' : 'Maon Moss',
#     '{COMPANY_NAME}' : 'Comprac Holdings (Pty) LTD',
#     '{TOTAL_MONTHS}' : '2',
#     '{MONTHLY_FEE}' : '',
#     '{REF}' : 'WC2215/0624',
#     '{COMPANY_TEL_NUMBER}' : '(000) 123 1234'
#     # '{INSPECTION_DAY}' : '',
#     # '{VISITS_PER_MONTH}' : '',
# }




def create_proposal(pdf, word, company):
    # print('CREATE PROPOSAL')

    output_path = 'output.docx'

    extracted_data = extract_data(pdf)

    match word:
        case 'utils/pc.docx':
            extra_directors = 'R. Looch (non-Executive)'
        case 'utils/pc_retainer.docx':
            extra_directors = 'R. Looch (non-Executive)'
        case 'utils/mobi_pack.docx':
            extra_directors = 'R. Looch (non-Executive)'

    placeholders = {}
    for d in list(extracted_data.keys()):
        key = '{' + d.upper() + '}'
        if d == 'client_address':
            placeholders[key] = extracted_data[d].replace(', ', '\n')
        elif d == 'inspection_day':
            placeholders[key] = inspection_day
        elif d == 'total_fee':
            total_fee = extracted_data[d].split('.')[0].replace(',', '')
            monthly_fee = int(total_fee) / int(extracted_data['total_months'])
            placeholders['{MONTHLY_FEE}'] = str(monthly_fee)
        else:
            placeholders[key] = extracted_data[d]
    placeholders['{COMPANY_NAME}'] = company
    placeholders['{REF}'] = 'WC2215/0624'
    placeholders['{COMPANY_TEL_NUMBER}'] = '(011) 425 6352'

    print('###', word)

    # Replace placeholders and save the DOCX
    output = read_proposal(word, output_path, placeholders)
    
    print("DOCX and PDF files have been saved.")
    
    return output







# EXTRACT & SAVE DATA FROM scope FORM FIELDS:

def extract_data(file):
    # Read the content of the InMemoryUploadedFile
    file_content = file.read()

    # Use BytesIO to create a file-like object from the content
    pdf_file = BytesIO(file_content)

    try:
        extracted_data = {}
        scope_reader = PdfReader(file)
        scope_number_of_pages = len(scope_reader.pages)
        scope_page = scope_reader.pages[0]
        scope_text = scope_page.extract_text()
        scope_fields = scope_reader.get_fields()

        for field in scope_fields:
            question = scope_fields[field]['/T']
            answer = scope_fields[field]['/V']
            extracted_data[question] = answer
            # print(f'{question}: {answer}')

        extracted_data['date'] = '15 July 2024'
        # if int(extracted_data['project_value']) >= 60000000:
        #     extracted_data['switch'] = permit
        # else:
        #     extracted_data['switch'] = ''
        # if extracted_data['inspection_day'] != '':
        #     inspection_day = f'Inspections will be conducted on a {extracted_data['inspection_day']}.'

        return extracted_data

    finally:
        pdf_file.close()






def read_proposal(doc_path, output_doc_path, placeholders):
# def read_proposal(doc_path, placeholders):
    # Load the document
    doc = Document(doc_path)


    # WORKING WITH HEADERS AND FOOTERS:
    
    for s in doc.sections:
        # Creating the first page legal footer:
        image_par = s.first_page_footer.add_paragraph()
        # fetch image from database:
        image_instance = fetch_image_from_database('proposal_footer_image')
        if image_instance.image:
            # image_bytes = image_instance.image.read()
            image_par.add_run().add_picture(image_instance.image, width=Cm(1))
        image_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph = s.first_page_footer.add_paragraph()
        for par in FOOTER_LEGAL:
            run = paragraph.add_run(par['text'])
            run.font.name = 'Arial Narrow'
            run.font.size = Pt(8)
            match par['formatting']:
                case 'red':
                    run.font.color.rgb = RGBColor(0xff, 0x00, 0x00)
                case 'small':
                    run.font.size = Pt(6)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        replace_placeholders(placeholders, paragraph, output_doc_path)

        for paragraph in s.first_page_header.paragraphs:
            # print(paragraph.text)
            replace_placeholders(placeholders, paragraph, output_doc_path)
        for paragraph in s.header.paragraphs:
            # print(paragraph.text)
            replace_placeholders(placeholders, paragraph, output_doc_path)


    
    # WORKING WITH TABLES:

    # add rows for total months:
    for m in range(1, int(placeholders['{TOTAL_MONTHS}']), 1):
        row = doc.tables[0].add_row()
        row.height = Cm(0.75)
        row.cells[0].add_paragraph().add_run(f'Month {m+1}').bold = True
        row.cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell_par = row.cells[1].add_paragraph()
        monthly_fee_run = cell_par.add_run('R {MONTHLY_FEE}')
        monthly_fee_run.font.name = 'Arial Narrow'
        monthly_fee_run.font.size = Pt(11)
        monthly_fee_run.font.color.rgb = RGBColor(0xff, 0x00, 0x00)
        cell_par.add_run(' per month exc. VAT').font.name = 'Arial Narrow'
        row.cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
    for table in doc.tables:
        # print()
        # print(table)
        for row in table.rows:
            # print(row)
            for cell in row.cells:
                # print(cell.paragraphs)
                for paragraph in cell.paragraphs:
                    replace_placeholders(placeholders, paragraph, output_doc_path)
        

    # WORKING WITH PARAGRAPHS:

    for paragraph in doc.paragraphs:
        # print(paragraph)
        replace_placeholders(placeholders, paragraph, output_doc_path)
        # replace_placeholders(placeholders, paragraph, 'test.docx')
        
    
    # doc.save('test.docx')
    doc.save(output_doc_path)
    return output_doc_path
    # return doc




    

    
    
    
    
    
    
    
    
    
# Replace placeholders in the document

def replace_placeholders(placeholders, paragraph, output_doc_path):
    
    # print()
    # print(paragraph)

    if len(paragraph.runs) > 0:
        # create runs list
        runs = [] # [{'text': [], 'formatting': (None, None, None)}]
        # create text_formatting list
        text_formatting = [] # [('text', (None, None, None))]
        current_run = 0
        current_formatting = 0
        # separate each text into text_formatting list with the style applied to that text in tuple
        for run in paragraph.runs:
            # print(f'\"{run.text}\"')
            # get run formatting:
            formatting = (run.bold, run.italic, run.font.color.rgb, run.font.name, run.font.size)
            text_formatting.append((run.text, formatting))
        # print(f'TF: {text_formatting}')

        # if run contains images, save those
        
        # clear all runs from paragraph to start anew:
        paragraph.clear()

        # if runs == []: add a new {text: [], formatting: ()}, set the formatting to the first run in runs, set current run
        if runs == []:
            current_run = 0
            runs.append({'text': [], 'formatting': text_formatting[0][1]})
            # print(runs[current_run]['formatting'])
        # print(runs)
        # read through text_formatting list and for each tuple:
        for text in text_formatting:
        # if this tuple has the same formatting as the current run, add the text to that run, ??remove that tuple from text_formatting list??:
            if text[1] == runs[current_run]['formatting']:
        # DEPRECATE THE FOLLOWING LINE:
        #         if not text[0].isspace() and text[0] != '':
                runs[current_run]['text'].append(text[0])
                current_formatting += 1
        # else, create a new run, add it to the runs list, set it as current run, add the text and formatting to that run, continue ...
            else:
                if not text[0].isspace() and text[0] != '':
                    runs.append({'text': [text[0]], 'formatting': text[1]})
                current_run += 1
        # print(f'RUNS: {runs}')

        # join text in each of the runs['text'] lists NOT?? with a space:
        for run in runs:
            run['text'] = ''.join(run['text'])
            # run['text'] = ' '.join(run['text'])
            # print(run['text'])

            # in this text, REPLACE placeholders with data, USING REGEXs:
            matches = re.findall(r'\{\w*\^?\}', run['text'])
            # print(f'MATCHES: {matches}')
            replaced = run['text']
            for m in matches:
                if '^' in m:
                    replaced = replaced.replace('^', '', 1)
                    m = m.replace('^', '', 1)
                    replaced = replaced.replace(m, placeholders[m].upper())
                else:
                    replaced = replaced.replace(m, placeholders[m])
            run['text'] = replaced

            # add this run to the paragraph as a new whole run
            new_run = paragraph.add_run(text=run['text'])
            new_run.bold = run['formatting'][0]
            new_run.italic = run['formatting'][1]
            new_run.font.color.rgb = run['formatting'][2]
            new_run.font.name = run['formatting'][3]
            new_run.font.size = run['formatting'][4]
    





# create_proposal("convert/utils/scope.pdf", "convert/utils/proposal.docx")
# read_proposal('convert/utils/prop.docx', 'test.docx', placeholders)







# HELPER FUNCTION:

def fetch_image_from_database(image_name):
    ...
    try:
        image_instance = Image.objects.get(name=image_name)
        return image_instance
    except Image.DoesNotExist:
        return None